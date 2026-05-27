#!/usr/bin/env python3
"""
security-mcp — Reverse-engineering MCP for malware and connected-device analysis.

Built 2026-05-21 for the Atoto/Suding investigation workflow. Tools are designed
for ANALYSIS of devices the operator owns and authorizes themselves to study,
not for ATTACKING third-party systems.

All tools accept absolute paths. All subprocesses bounded by timeout. Output is
truncated above MAX_OUTPUT_BYTES to keep the LLM context manageable.

Run with: python3 server.py
Or via the run.sh launcher.
"""

from __future__ import annotations

import asyncio
import hashlib
import json
import os
import re
import shlex
import shutil
import sqlite3
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Any

from mcp.server.fastmcp import FastMCP

# Configuration --------------------------------------------------------------

DEFAULT_TIMEOUT = 60
LONG_TIMEOUT = 600
MAX_OUTPUT_BYTES = 200 * 1024  # 200 KB cap on any single tool return

SCRATCH = Path.home() / "security-mcp" / "scratch"
SCRATCH.mkdir(parents=True, exist_ok=True)

# Helpers --------------------------------------------------------------------

def _truncate(text: str, limit: int = MAX_OUTPUT_BYTES) -> str:
    if len(text) <= limit:
        return text
    head = limit // 2
    tail = limit - head - 200
    return text[:head] + f"\n\n... [truncated {len(text)-limit} bytes; showing {head} head + {tail} tail] ...\n\n" + text[-tail:]


def _run(argv: list[str], timeout: int = DEFAULT_TIMEOUT, capture_stderr: bool = True, input_bytes: bytes | None = None, cwd: str | None = None) -> dict[str, Any]:
    """Run a subprocess and return {stdout, stderr, returncode, command}. Bounded by timeout."""
    try:
        proc = subprocess.run(
            argv,
            timeout=timeout,
            capture_output=True,
            input=input_bytes,
            cwd=cwd,
        )
        return {
            "command": " ".join(shlex.quote(a) for a in argv),
            "returncode": proc.returncode,
            "stdout": _truncate(proc.stdout.decode("utf-8", errors="replace")),
            "stderr": _truncate(proc.stderr.decode("utf-8", errors="replace")) if capture_stderr else "",
        }
    except subprocess.TimeoutExpired:
        return {"command": " ".join(shlex.quote(a) for a in argv), "returncode": -1, "stdout": "", "stderr": f"TIMEOUT after {timeout}s"}
    except FileNotFoundError as e:
        return {"command": " ".join(shlex.quote(a) for a in argv), "returncode": -1, "stdout": "", "stderr": f"FileNotFoundError: {e}"}
    except Exception as e:
        return {"command": " ".join(shlex.quote(a) for a in argv), "returncode": -1, "stdout": "", "stderr": f"{type(e).__name__}: {e}"}


def _require_path(path: str, must_exist: bool = True) -> Path:
    p = Path(path).expanduser().resolve()
    if must_exist and not p.exists():
        raise FileNotFoundError(f"Path does not exist: {p}")
    return p


# MCP server -----------------------------------------------------------------

mcp = FastMCP(
    name="security-mcp",
    instructions=(
        "Reverse-engineering toolkit for malware and connected-device analysis. "
        "Operates on local files only. All paths must be absolute. "
        "Subprocesses are timeout-bounded. "
        "Designed for analysis of hardware the operator owns and has authorization to study."
    ),
)


# ----- File / binary utilities ------------------------------------------------

@mcp.tool()
def file_info(path: str) -> dict:
    """Identify a file: size, magic, MIME type, basic structure. Wraps `file(1)` and adds size."""
    p = _require_path(path)
    info = _run(["file", "--brief", "--mime", str(p)])
    info2 = _run(["file", "--brief", str(p)])
    return {
        "path": str(p),
        "size_bytes": p.stat().st_size,
        "mime": info["stdout"].strip(),
        "describe": info2["stdout"].strip(),
    }


@mcp.tool()
def sha256(path: str) -> dict:
    """Compute SHA-256 of a single file."""
    p = _require_path(path)
    h = hashlib.sha256()
    with p.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1 << 20), b""):
            h.update(chunk)
    return {"path": str(p), "sha256": h.hexdigest(), "size_bytes": p.stat().st_size}


@mcp.tool()
def sha256_manifest(directory: str, recursive: bool = True, save_to: str | None = None) -> dict:
    """SHA-256 every file under a directory. Optionally write the manifest to a file. Returns the manifest as a string."""
    d = _require_path(directory)
    if not d.is_dir():
        raise NotADirectoryError(d)
    lines = []
    walker = d.rglob("*") if recursive else d.glob("*")
    for p in sorted(walker):
        if p.is_file():
            h = hashlib.sha256()
            try:
                with p.open("rb") as fh:
                    for chunk in iter(lambda: fh.read(1 << 20), b""):
                        h.update(chunk)
                rel = p.relative_to(d)
                lines.append(f"{h.hexdigest()}  ./{rel}")
            except (PermissionError, OSError) as e:
                lines.append(f"# ERROR {p}: {e}")
    manifest = "\n".join(lines) + "\n"
    out = {"directory": str(d), "files_hashed": len(lines), "manifest_preview": _truncate(manifest, 8000)}
    if save_to:
        sp = Path(save_to).expanduser()
        sp.write_text(manifest)
        out["written_to"] = str(sp)
    return out


@mcp.tool()
def hex_dump(path: str, offset: int = 0, length: int = 512) -> str:
    """Hexdump bytes from a file. Wraps `xxd` with -s offset -l length."""
    p = _require_path(path)
    return _truncate(_run(["xxd", "-s", str(offset), "-l", str(length), str(p)])["stdout"])


@mcp.tool()
def binary_strings(path: str, min_len: int = 4, encoding: str = "s", grep_pattern: str | None = None, limit: int = 5000) -> str:
    """Extract printable strings from a binary. Optionally grep for a regex. encoding: s=7bit/L,b,l,B,S=16/32bit."""
    p = _require_path(path)
    argv = ["strings", "-a", "-n", str(min_len), "-e", encoding, str(p)]
    out = _run(argv, timeout=LONG_TIMEOUT)["stdout"]
    if grep_pattern:
        regex = re.compile(grep_pattern, re.IGNORECASE)
        lines = [ln for ln in out.splitlines() if regex.search(ln)]
    else:
        lines = out.splitlines()
    return _truncate("\n".join(lines[:limit]))


@mcp.tool()
def elf_info(path: str) -> dict:
    """ELF file information: header, dynamic deps, symbols. Wraps `readelf` + `nm`."""
    p = _require_path(path)
    return {
        "header": _truncate(_run(["readelf", "-h", str(p)])["stdout"], 4000),
        "dynamic_deps": _truncate(_run(["readelf", "-d", str(p)])["stdout"], 8000),
        "symbols": _truncate(_run(["nm", "-D", "--defined-only", str(p)])["stdout"], 20000),
        "sections": _truncate(_run(["readelf", "-S", str(p)])["stdout"], 8000),
    }


# ----- Android / APK ---------------------------------------------------------

@mcp.tool()
def apk_manifest(apk_path: str) -> dict:
    """Dump APK manifest: package, permissions, activities, services, providers."""
    p = _require_path(apk_path)
    badging = _run(["aapt", "dump", "badging", str(p)], timeout=120)["stdout"]
    perms = _run(["aapt", "dump", "permissions", str(p)], timeout=60)["stdout"]
    return {
        "badging": _truncate(badging, 20000),
        "permissions": _truncate(perms, 10000),
    }


@mcp.tool()
def apk_strings(apk_path: str, grep_pattern: str | None = None, dex_only: bool = True, limit: int = 5000) -> str:
    """Extract strings from an APK's classes.dex files. Optionally regex-grep."""
    p = _require_path(apk_path)
    with tempfile.TemporaryDirectory(dir=str(SCRATCH)) as td:
        _run(["unzip", "-q", "-o", str(p), "-d", td], timeout=120)
        dexes = sorted(Path(td).glob("classes*.dex"))
        if not dex_only:
            extra = [x for x in Path(td).rglob("*") if x.is_file() and x.suffix in {".so", ".bin"}]
            dexes.extend(extra)
        all_lines = []
        for dex in dexes:
            out = _run(["strings", "-a", "-n", "4", str(dex)], timeout=120)["stdout"]
            all_lines.extend(out.splitlines())
    if grep_pattern:
        regex = re.compile(grep_pattern, re.IGNORECASE)
        all_lines = [ln for ln in all_lines if regex.search(ln)]
    return _truncate("\n".join(all_lines[:limit]))


@mcp.tool()
def apk_signature(apk_path: str) -> dict:
    """Inspect APK signing certs via apksigner verify or unzip+openssl on META-INF."""
    p = _require_path(apk_path)
    apksigner = _run(["apksigner", "verify", "--print-certs", str(p)], timeout=60)
    if apksigner["returncode"] == 0:
        return {"tool": "apksigner", "output": _truncate(apksigner["stdout"], 10000)}
    with tempfile.TemporaryDirectory(dir=str(SCRATCH)) as td:
        _run(["unzip", "-q", "-o", str(p), "META-INF/*", "-d", td], timeout=30)
        certs = sorted(Path(td).glob("META-INF/*.RSA")) + sorted(Path(td).glob("META-INF/*.DSA")) + sorted(Path(td).glob("META-INF/*.EC"))
        out = []
        for c in certs:
            res = _run(["openssl", "pkcs7", "-inform", "DER", "-print_certs", "-text", "-in", str(c)], timeout=30)
            out.append(f"--- {c.name} ---\n{res['stdout']}")
        return {"tool": "openssl-on-META-INF", "output": _truncate("\n".join(out), 20000)}


@mcp.tool()
def apk_decompile(apk_path: str, out_dir: str | None = None, no_res: bool = True) -> dict:
    """Decompile an APK to Java sources via jadx. Returns the output directory path."""
    p = _require_path(apk_path)
    jadx = shutil.which("jadx") or "/tmp/jadx/bin/jadx"
    if not Path(jadx).exists():
        return {"error": "jadx not found; install or place at /tmp/jadx/bin/jadx"}
    if not out_dir:
        out_dir = str(SCRATCH / f"jadx-{p.stem}")
    out = Path(out_dir).expanduser()
    out.mkdir(parents=True, exist_ok=True)
    argv = [jadx, "-d", str(out), str(p)]
    if no_res:
        argv.append("--no-res")
    res = _run(argv, timeout=LONG_TIMEOUT)
    return {"out_dir": str(out), "returncode": res["returncode"], "log_tail": _truncate(res["stderr"], 4000)}


@mcp.tool()
def search_decompiled(directory: str, pattern: str, file_pattern: str = "*.java", max_results: int = 200) -> str:
    """Grep recursively under a decompiled APK directory. Pattern is a regex."""
    d = _require_path(directory)
    out_lines = []
    rg = shutil.which("rg")
    if rg:
        res = _run([rg, "--no-heading", "-n", "-g", file_pattern, "-e", pattern, str(d)], timeout=120)
        return _truncate(res["stdout"])
    # Fallback to grep -rn
    res = _run(["grep", "-rn", "--include", file_pattern, "-E", pattern, str(d)], timeout=180)
    return _truncate(res["stdout"])


# ----- Android filesystem images --------------------------------------------

@mcp.tool()
def lpunpack_super(super_img: str, partition: str | None = None, out_dir: str | None = None, info_only: bool = False) -> dict:
    """Extract logical partitions from a super.img using a vendored lpunpack.py.
    If partition is None and info_only=True, prints the partition table.
    """
    s = _require_path(super_img)
    lpunpack = Path("/tmp/lpunpack/lpunpack.py")
    if not lpunpack.exists():
        return {"error": "lpunpack.py not found at /tmp/lpunpack/lpunpack.py — download from github.com/unix3dgforce/lpunpack"}
    if info_only:
        res = _run(["python3", str(lpunpack), "--info", str(s)], timeout=LONG_TIMEOUT)
        return {"info": _truncate(res["stdout"])}
    if not out_dir:
        out_dir = str(SCRATCH / f"lpunpack-{s.stem}")
    o = Path(out_dir).expanduser()
    o.mkdir(parents=True, exist_ok=True)
    argv = ["python3", str(lpunpack), str(s), str(o)]
    if partition:
        argv = ["python3", str(lpunpack), "-p", partition, str(s), str(o)]
    res = _run(argv, timeout=LONG_TIMEOUT * 2)
    return {"out_dir": str(o), "returncode": res["returncode"], "log_tail": _truncate(res["stdout"] + res["stderr"], 4000)}


@mcp.tool()
def mount_partition_image(img_path: str, mount_point: str, options: str = "ro,loop,norecovery", sudo: bool = True) -> dict:
    """Loop-mount an ext4 partition image read-only. Returns the mount status."""
    p = _require_path(img_path)
    mp = Path(mount_point).expanduser()
    mp.mkdir(parents=True, exist_ok=True)
    argv = (["sudo"] if sudo else []) + ["mount", "-o", options, str(p), str(mp)]
    res = _run(argv, timeout=30)
    return {"mounted_at": str(mp), "image": str(p), "returncode": res["returncode"], "stderr": res["stderr"]}


@mcp.tool()
def unmount(mount_point: str, sudo: bool = True) -> dict:
    """Unmount a mountpoint. Use sudo if needed."""
    argv = (["sudo"] if sudo else []) + ["umount", mount_point]
    res = _run(argv, timeout=30)
    return res


# ----- SQLite / Android data ------------------------------------------------

@mcp.tool()
def sqlite_tables(db_path: str) -> dict:
    """List tables and basic row counts in a SQLite database."""
    p = _require_path(db_path)
    conn = sqlite3.connect(f"file:{p}?mode=ro", uri=True)
    try:
        cur = conn.cursor()
        tables = [r[0] for r in cur.execute("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name").fetchall()]
        counts = {}
        for t in tables:
            try:
                counts[t] = cur.execute(f'SELECT COUNT(*) FROM "{t}"').fetchone()[0]
            except sqlite3.DatabaseError as e:
                counts[t] = f"err: {e}"
        return {"db": str(p), "tables": counts}
    finally:
        conn.close()


@mcp.tool()
def sqlite_schema(db_path: str, table: str | None = None) -> str:
    """Show CREATE statements for one or all tables."""
    p = _require_path(db_path)
    conn = sqlite3.connect(f"file:{p}?mode=ro", uri=True)
    try:
        cur = conn.cursor()
        if table:
            row = cur.execute("SELECT sql FROM sqlite_master WHERE name=?", (table,)).fetchone()
            return row[0] if row else f"(no table {table!r})"
        rows = cur.execute("SELECT sql FROM sqlite_master WHERE type='table' AND sql IS NOT NULL ORDER BY name").fetchall()
        return "\n\n".join(r[0] for r in rows)
    finally:
        conn.close()


@mcp.tool()
def sqlite_query(db_path: str, sql: str, params: list | None = None, limit: int = 200) -> dict:
    """Run a SELECT against a SQLite database (read-only). Returns rows as dicts.
    Refuses statements that aren't SELECT / EXPLAIN / WITH / PRAGMA."""
    p = _require_path(db_path)
    s = sql.strip().lstrip("(").lower()
    if not (s.startswith("select") or s.startswith("explain") or s.startswith("with") or s.startswith("pragma")):
        return {"error": "Only read-only queries (SELECT/EXPLAIN/WITH/PRAGMA) are permitted."}
    conn = sqlite3.connect(f"file:{p}?mode=ro", uri=True)
    conn.row_factory = sqlite3.Row
    try:
        rows = [dict(r) for r in conn.execute(sql, params or []).fetchmany(limit)]
        return {"db": str(p), "sql": sql, "rows": rows, "row_count_capped_at": limit}
    except sqlite3.Error as e:
        return {"error": f"sqlite error: {e}", "sql": sql}
    finally:
        conn.close()


# ----- Android-specific parsers ---------------------------------------------

@mcp.tool()
def android_shared_prefs(xml_path: str) -> dict:
    """Parse an Android SharedPreferences XML file into a dict."""
    import xml.etree.ElementTree as ET
    p = _require_path(xml_path)
    tree = ET.parse(p)
    root = tree.getroot()
    out: dict[str, Any] = {}
    for el in root:
        name = el.attrib.get("name")
        if el.tag == "string":
            out[name] = el.text
        elif el.tag == "boolean":
            out[name] = el.attrib.get("value") == "true"
        elif el.tag in ("int", "long", "float"):
            out[name] = el.attrib.get("value")
        elif el.tag == "set":
            out[name] = [c.text for c in el]
        else:
            out[name] = ET.tostring(el, encoding="unicode")
    return {"path": str(p), "values": out}


@mcp.tool()
def find_setuid_setgid(directory: str) -> dict:
    """List all setuid / setgid binaries under a directory tree. Uses sudo find -perm."""
    d = _require_path(directory)
    res = _run(["sudo", "find", str(d), "-type", "f", "(", "-perm", "-u+s", "-o", "-perm", "-g+s", ")", "-printf", "%M %u %g %p\\n"], timeout=300)
    return {"directory": str(d), "results": _truncate(res["stdout"])}


@mcp.tool()
def parse_init_rc(init_rc_path: str) -> dict:
    """Parse an Android init.rc file. Extracts service definitions, on triggers, imports, and setprops."""
    p = _require_path(init_rc_path)
    text = p.read_text(errors="replace")
    services = []
    triggers = []
    imports = []
    setprops = []
    current_service = None
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("service "):
            parts = stripped.split(maxsplit=2)
            if len(parts) >= 3:
                current_service = {"name": parts[1], "exec": parts[2], "options": []}
                services.append(current_service)
            continue
        if stripped.startswith("on "):
            triggers.append(stripped)
            current_service = None
            continue
        if stripped.startswith("import "):
            imports.append(stripped)
            continue
        if "setprop" in stripped:
            setprops.append(stripped)
        if current_service and stripped and not stripped.startswith("#"):
            current_service["options"].append(stripped)
    return {
        "path": str(p),
        "service_count": len(services),
        "services": services[:200],
        "trigger_count": len(triggers),
        "triggers": triggers[:200],
        "imports": imports,
        "setprops": setprops[:200],
    }


@mcp.tool()
def android_url_extraction(directory: str, extensions: list[str] | None = None, limit: int = 500) -> str:
    """Find URLs and IP addresses in files under a directory (good for endpoint catalog from decompiled apps)."""
    d = _require_path(directory)
    exts = extensions or [".java", ".smali", ".xml", ".properties", ".so", ".bin"]
    pattern = r"https?://[A-Za-z0-9._~:/?#\[\]@!$&'()*+,;=%-]+|(?:\d{1,3}\.){3}\d{1,3}(?::\d{1,5})?"
    found = set()
    for p in d.rglob("*"):
        if p.is_file() and (not exts or p.suffix.lower() in exts):
            try:
                t = p.read_text(errors="replace")
                for m in re.findall(pattern, t):
                    found.add(m)
            except (UnicodeDecodeError, OSError):
                pass
            if len(found) > limit:
                break
    return _truncate("\n".join(sorted(found)[:limit]))


# ----- Subprocess escape hatch ----------------------------------------------

@mcp.tool()
def run_shell(command: str, timeout: int = 60, cwd: str | None = None) -> dict:
    """Run an arbitrary shell command via /bin/sh -c. Bounded by timeout. Use sparingly when no
    dedicated tool fits. Output is captured and truncated to MAX_OUTPUT_BYTES."""
    res = _run(["/bin/sh", "-c", command], timeout=timeout, cwd=cwd)
    return res


@mcp.tool()
def gpg_verify(signed_file: str, signature_file: str | None = None) -> dict:
    """Verify a GPG signature. If signature_file omitted, assume detached at signed_file+'.asc' or signed_file+'.sig'."""
    p = _require_path(signed_file)
    if signature_file is None:
        for ext in (".asc", ".sig"):
            cand = Path(str(p) + ext)
            if cand.exists():
                signature_file = str(cand)
                break
        if signature_file is None:
            return {"error": "no signature file found alongside (tried .asc and .sig)"}
    res = _run(["gpg", "--verify", signature_file, str(p)], timeout=30)
    return res


# ----- Network / endpoint helpers -------------------------------------------

@mcp.tool()
def whois_lookup(domain_or_ip: str) -> str:
    """Simple WHOIS lookup. Wraps `whois`."""
    if not shutil.which("whois"):
        return "whois binary not installed"
    return _truncate(_run(["whois", domain_or_ip], timeout=30)["stdout"], 12000)


# ===========================================================================
# Kali Linux / REMnux integrations
# ===========================================================================
# Tools wrapped here are for analysis, authorized testing, and research on
# hardware/networks the operator owns or has explicit authorization to study.
# The run_shell escape hatch covers anything not given a dedicated wrapper.
# ===========================================================================

def _have(binary: str) -> bool:
    return shutil.which(binary) is not None


# ----- Network reconnaissance -----------------------------------------------

@mcp.tool()
def nmap_scan(target: str, ports: str = "1-1000", scan_type: str = "S", extra_args: list[str] | None = None, timeout: int = 300) -> str:
    """Port scan via nmap. scan_type: S=SYN(needs root), T=connect, U=UDP, A=aggressive.
    Operator is responsible for ensuring target is owned or authorized."""
    if not _have("nmap"):
        return "nmap not installed (apt install nmap)"
    argv = ["sudo", "nmap"] if scan_type in ("S", "U", "A") else ["nmap"]
    argv += [f"-s{scan_type}", "-p", ports, target]
    if extra_args:
        argv += extra_args
    return _truncate(_run(argv, timeout=timeout)["stdout"])


@mcp.tool()
def nmap_version_scan(target: str, ports: str = "1-10000", timeout: int = 600) -> str:
    """Service / version detection via nmap -sV. Authorized targets only."""
    if not _have("nmap"):
        return "nmap not installed"
    return _truncate(_run(["sudo", "nmap", "-sS", "-sV", "-p", ports, target], timeout=timeout)["stdout"])


@mcp.tool()
def nmap_script_scan(target: str, script: str = "default", ports: str = "1-1000", timeout: int = 600) -> str:
    """Run nmap NSE scripts. script: e.g. 'default', 'vuln', 'http-enum', 'ssl-cert'."""
    if not _have("nmap"):
        return "nmap not installed"
    return _truncate(_run(["sudo", "nmap", "--script", script, "-p", ports, target], timeout=timeout)["stdout"])


@mcp.tool()
def masscan_scan(target: str, ports: str = "0-65535", rate: int = 1000, timeout: int = 600) -> str:
    """Fast wide-port scan via masscan. Requires root. Authorized targets only.
    Rate above 10000 can saturate links — start low."""
    if not _have("masscan"):
        return "masscan not installed"
    return _truncate(_run(["sudo", "masscan", "-p", ports, "--rate", str(rate), target], timeout=timeout)["stdout"])


@mcp.tool()
def dig_query(name: str, record_type: str = "A", server: str | None = None) -> str:
    """DNS lookup via dig. record_type: A, AAAA, MX, NS, TXT, SOA, ANY, etc."""
    if not _have("dig"):
        return "dig not installed (apt install dnsutils)"
    argv = ["dig", "+noall", "+answer", "+authority"]
    if server:
        argv.append(f"@{server}")
    argv += [name, record_type]
    return _truncate(_run(argv, timeout=30)["stdout"])


@mcp.tool()
def traceroute_to(target: str, max_hops: int = 30) -> str:
    """traceroute to target."""
    if not _have("traceroute"):
        return "traceroute not installed"
    return _truncate(_run(["traceroute", "-m", str(max_hops), target], timeout=120)["stdout"])


@mcp.tool()
def httpx_probe(targets: list[str], timeout: int = 60) -> str:
    """HTTP probe via projectdiscovery httpx — fast check of which targets respond on HTTP/S."""
    if not _have("httpx"):
        return "httpx not installed"
    with tempfile.NamedTemporaryFile("w", suffix=".txt", delete=False, dir=str(SCRATCH)) as f:
        f.write("\n".join(targets))
        target_file = f.name
    try:
        return _truncate(_run(["httpx", "-silent", "-status-code", "-title", "-tech-detect", "-l", target_file], timeout=timeout)["stdout"])
    finally:
        Path(target_file).unlink(missing_ok=True)


@mcp.tool()
def theharvester_emails(domain: str, source: str = "bing", timeout: int = 120) -> str:
    """Passive email / subdomain harvest. source: bing, google, duckduckgo, virustotal, certspotter, etc."""
    if not _have("theHarvester") and not _have("theharvester"):
        return "theharvester not installed"
    bin_ = shutil.which("theHarvester") or shutil.which("theharvester")
    return _truncate(_run([bin_, "-d", domain, "-b", source, "-l", "200"], timeout=timeout)["stdout"])


# ----- Web app probing ------------------------------------------------------

@mcp.tool()
def nikto_scan(url: str, timeout: int = 300) -> str:
    """Web vulnerability scanner. Authorized targets only."""
    if not _have("nikto"):
        return "nikto not installed"
    return _truncate(_run(["nikto", "-h", url, "-ask", "no"], timeout=timeout)["stdout"])


@mcp.tool()
def whatweb_fingerprint(url: str, aggression: int = 1) -> str:
    """Identify web tech stack. aggression 1=stealthy, 3=aggressive, 4=heavy."""
    if not _have("whatweb"):
        return "whatweb not installed"
    return _truncate(_run(["whatweb", f"-a{aggression}", "--no-errors", url], timeout=60)["stdout"])


@mcp.tool()
def gobuster_dir(url: str, wordlist: str = "/usr/share/wordlists/dirb/common.txt", threads: int = 10, timeout: int = 300) -> str:
    """Directory brute-force via gobuster. Authorized targets only."""
    if not _have("gobuster"):
        return "gobuster not installed"
    return _truncate(_run(["gobuster", "dir", "-u", url, "-w", wordlist, "-t", str(threads), "-q"], timeout=timeout)["stdout"])


@mcp.tool()
def ffuf_fuzz(url_template: str, wordlist: str = "/usr/share/wordlists/dirb/common.txt", match_codes: str = "200,301,302,403", timeout: int = 300) -> str:
    """Fuzzing via ffuf. url_template uses FUZZ as the substitution placeholder, e.g. https://x/FUZZ."""
    if not _have("ffuf"):
        return "ffuf not installed"
    return _truncate(_run(["ffuf", "-u", url_template, "-w", wordlist, "-mc", match_codes, "-s"], timeout=timeout)["stdout"])


@mcp.tool()
def sqlmap_test(url: str, data: str | None = None, cookie: str | None = None, level: int = 1, risk: int = 1, timeout: int = 600, owner_asserts_authorized: bool = False) -> str:
    """Automated SQL injection probing. The operator must own the target or have written authorization.
    Set owner_asserts_authorized=True after confirming authorization."""
    if not owner_asserts_authorized:
        return "sqlmap not invoked. Set owner_asserts_authorized=True after confirming the target is yours or you have written pen-test authorization."
    if not _have("sqlmap"):
        return "sqlmap not installed"
    argv = ["sqlmap", "-u", url, "--batch", "--level", str(level), "--risk", str(risk)]
    if data:
        argv += ["--data", data]
    if cookie:
        argv += ["--cookie", cookie]
    return _truncate(_run(argv, timeout=timeout)["stdout"])


@mcp.tool()
def curl_request(url: str, method: str = "GET", headers: dict | None = None, body: str | None = None, follow_redirects: bool = True, timeout: int = 60) -> dict:
    """HTTP request via curl with full control over method, headers, body. Returns headers + body."""
    argv = ["curl", "-sS", "-D", "-", "--max-time", str(timeout), "-X", method, url]
    if follow_redirects:
        argv.append("-L")
    if headers:
        for k, v in headers.items():
            argv += ["-H", f"{k}: {v}"]
    if body:
        argv += ["--data-binary", body]
    res = _run(argv, timeout=timeout + 5)
    out = res["stdout"]
    # Split headers/body on the first blank line(s)
    parts = re.split(r"\r?\n\r?\n", out, maxsplit=1)
    return {
        "url": url,
        "headers": parts[0] if parts else "",
        "body": _truncate(parts[1], 100000) if len(parts) > 1 else "",
        "stderr": res["stderr"][:1000] if res["stderr"] else "",
    }


# ----- TLS / certificate inspection -----------------------------------------

@mcp.tool()
def openssl_s_client(host: str, port: int = 443, sni: str | None = None, timeout: int = 30) -> str:
    """Fetch TLS server certificate via openssl s_client. Returns the negotiated cert and details."""
    sni_arg = sni or host
    cmd = f"echo | openssl s_client -connect {shlex.quote(host)}:{port} -servername {shlex.quote(sni_arg)} -showcerts 2>/dev/null | openssl x509 -text -noout"
    return _truncate(_run(["/bin/sh", "-c", cmd], timeout=timeout)["stdout"])


@mcp.tool()
def openssl_x509_decode(cert_file: str) -> str:
    """Decode an X.509 cert file (PEM or DER) into human-readable form."""
    p = _require_path(cert_file)
    res = _run(["openssl", "x509", "-in", str(p), "-text", "-noout"], timeout=30)
    if res["returncode"] == 0:
        return _truncate(res["stdout"])
    return _truncate(_run(["openssl", "x509", "-inform", "DER", "-in", str(p), "-text", "-noout"], timeout=30)["stdout"])


@mcp.tool()
def sslscan_target(host: str, port: int = 443) -> str:
    """Test SSL/TLS configuration of a target (ciphers, protocols, vulns)."""
    if not _have("sslscan"):
        return "sslscan not installed (apt install sslscan)"
    return _truncate(_run(["sslscan", "--no-colour", f"{host}:{port}"], timeout=120)["stdout"])


# ----- Forensics ------------------------------------------------------------

@mcp.tool()
def exiftool_metadata(path: str, recursive: bool = False) -> str:
    """Extract all metadata from a file (or recursively from a dir) via exiftool."""
    if not _have("exiftool"):
        return "exiftool not installed (apt install libimage-exiftool-perl)"
    argv = ["exiftool", "-a", "-u", "-g1"]
    if recursive:
        argv += ["-r", path]
    else:
        argv.append(path)
    return _truncate(_run(argv, timeout=120)["stdout"], 50000)


@mcp.tool()
def foremost_carve(disk_image: str, output_dir: str, types: str = "all") -> dict:
    """File carving via foremost. types: comma-separated extensions, or 'all'."""
    if not _have("foremost"):
        return {"error": "foremost not installed"}
    out = Path(output_dir).expanduser()
    out.mkdir(parents=True, exist_ok=True)
    argv = ["foremost", "-t", types, "-i", disk_image, "-o", str(out)]
    res = _run(argv, timeout=LONG_TIMEOUT)
    return {"out_dir": str(out), "returncode": res["returncode"], "stdout_tail": _truncate(res["stdout"], 8000)}


@mcp.tool()
def binwalk_analyze(file_path: str, options: str = "") -> str:
    """Analyze a binary/firmware file for embedded files and signatures."""
    if not _have("binwalk"):
        return "binwalk not installed"
    argv = ["binwalk", file_path]
    if options:
        argv += options.split()
    return _truncate(_run(argv, timeout=300)["stdout"])


@mcp.tool()
def binwalk_extract(file_path: str, output_dir: str | None = None) -> dict:
    """Extract embedded files from a firmware image via binwalk -eM (recursive)."""
    if not _have("binwalk"):
        return {"error": "binwalk not installed"}
    if not output_dir:
        output_dir = str(SCRATCH / f"binwalk-{Path(file_path).stem}")
    out = Path(output_dir).expanduser()
    out.mkdir(parents=True, exist_ok=True)
    argv = ["binwalk", "-eM", "-C", str(out), file_path]
    res = _run(argv, timeout=LONG_TIMEOUT * 2)
    return {"out_dir": str(out), "returncode": res["returncode"], "log": _truncate(res["stdout"], 8000)}


@mcp.tool()
def fls_list(image_path: str, offset: int | None = None) -> str:
    """List files/directories in a forensic image via Sleuthkit fls."""
    if not _have("fls"):
        return "fls not installed (apt install sleuthkit)"
    argv = ["fls", "-r"]
    if offset is not None:
        argv += ["-o", str(offset)]
    argv.append(image_path)
    return _truncate(_run(argv, timeout=300)["stdout"], 100000)


# ----- Malware / document analysis ------------------------------------------

@mcp.tool()
def yara_scan(rules_file: str, target_path: str, recursive: bool = False) -> str:
    """Scan a file or directory with YARA rules. Returns matching files + rule names."""
    if not _have("yara"):
        return "yara not installed (apt install yara)"
    argv = ["yara"]
    if recursive:
        argv.append("-r")
    argv += [rules_file, target_path]
    return _truncate(_run(argv, timeout=300)["stdout"])


@mcp.tool()
def ssdeep_hash(path: str) -> dict:
    """Compute ssdeep fuzzy hash of a file."""
    if not _have("ssdeep"):
        return {"error": "ssdeep not installed"}
    res = _run(["ssdeep", "-c", "-b", path], timeout=60)
    return {"path": path, "ssdeep_csv": res["stdout"]}


@mcp.tool()
def ssdeep_compare(file1: str, file2: str) -> str:
    """Compare two files by ssdeep similarity (0-100, higher = more similar)."""
    if not _have("ssdeep"):
        return "ssdeep not installed"
    return _truncate(_run(["ssdeep", "-d", file1, file2], timeout=60)["stdout"])


@mcp.tool()
def pdf_extract_text(pdf_path: str) -> str:
    """Extract plaintext from a PDF (pdftotext) — useful for analyzing suspicious PDFs without rendering."""
    if not _have("pdftotext"):
        return "pdftotext not installed (apt install poppler-utils)"
    p = _require_path(pdf_path)
    return _truncate(_run(["pdftotext", str(p), "-"], timeout=60)["stdout"])


# ----- Password / hash tools ------------------------------------------------

@mcp.tool()
def john_crack(hash_file: str, wordlist: str | None = None, format_: str | None = None, timeout: int = 600, owner_asserts_authorized: bool = False) -> str:
    """Crack hashes via John. The operator must own the hashes or have authorization to crack them.
    Set owner_asserts_authorized=True after confirming."""
    if not owner_asserts_authorized:
        return "john not invoked. Set owner_asserts_authorized=True after confirming you own the hashes."
    if not _have("john"):
        return "john not installed"
    argv = ["john"]
    if wordlist:
        argv += [f"--wordlist={wordlist}"]
    if format_:
        argv += [f"--format={format_}"]
    argv.append(hash_file)
    return _truncate(_run(argv, timeout=timeout)["stdout"])


@mcp.tool()
def hashcat_crack(hash_file: str, hashtype: str, wordlist: str, mode: int = 0, timeout: int = 600, owner_asserts_authorized: bool = False) -> str:
    """Crack hashes via hashcat. hashtype is the hashcat -m number. The operator must own the hashes."""
    if not owner_asserts_authorized:
        return "hashcat not invoked. Set owner_asserts_authorized=True after confirming you own the hashes."
    if not _have("hashcat"):
        return "hashcat not installed"
    argv = ["hashcat", "-m", str(hashtype), "-a", str(mode), hash_file, wordlist]
    return _truncate(_run(argv, timeout=timeout)["stdout"])


# ----- Network capture / packet analysis ------------------------------------

@mcp.tool()
def tcpdump_capture(interface: str, filter_expr: str = "", count: int = 100, output_pcap: str | None = None, timeout: int = 60) -> dict:
    """Capture packets via tcpdump. Bounded by count and timeout. output_pcap saves to file if given."""
    if not _have("tcpdump"):
        return {"error": "tcpdump not installed"}
    argv = ["sudo", "tcpdump", "-i", interface, "-c", str(count), "-nn"]
    if output_pcap:
        argv += ["-w", output_pcap]
    if filter_expr:
        argv += shlex.split(filter_expr)
    res = _run(argv, timeout=timeout)
    return {"interface": interface, "filter": filter_expr, "stdout": _truncate(res["stdout"], 50000), "stderr": _truncate(res["stderr"], 4000)}


@mcp.tool()
def tshark_analyze(pcap_path: str, display_filter: str | None = None, fields: list[str] | None = None, count: int = 200) -> str:
    """Read a pcap file via tshark. fields: e.g. ['frame.number','ip.src','tcp.dstport','http.host']."""
    if not _have("tshark"):
        return "tshark not installed"
    argv = ["tshark", "-r", pcap_path, "-c", str(count)]
    if display_filter:
        argv += ["-Y", display_filter]
    if fields:
        argv += ["-T", "fields"] + [a for f in fields for a in ("-e", f)]
    return _truncate(_run(argv, timeout=300)["stdout"])


@mcp.tool()
def tshark_summary(pcap_path: str) -> str:
    """Summarize a pcap: protocol hierarchy + endpoint counts."""
    if not _have("tshark"):
        return "tshark not installed"
    s1 = _run(["tshark", "-r", pcap_path, "-q", "-z", "io,phs"], timeout=300)["stdout"]
    s2 = _run(["tshark", "-r", pcap_path, "-q", "-z", "endpoints,ip"], timeout=300)["stdout"]
    return _truncate("=== Protocol hierarchy ===\n" + s1 + "\n=== Endpoints ===\n" + s2)


# ----- Discovery / introspection --------------------------------------------

@mcp.tool()
def discover_kali_remnux_tools() -> dict:
    """Probe which Kali/REMnux tools are available on PATH. Useful before invoking a tool."""
    categories = {
        "network_recon": ["nmap", "masscan", "rustscan", "dig", "nslookup", "whois", "host", "traceroute", "mtr", "httpx", "subfinder", "amass", "theHarvester", "theharvester"],
        "web_app": ["nikto", "whatweb", "gobuster", "ffuf", "wpscan", "sqlmap", "dirb", "wfuzz", "nuclei", "curl", "wget", "wpscan"],
        "packet_capture": ["tcpdump", "tshark", "wireshark", "mitmproxy", "mitmdump"],
        "tls": ["openssl", "sslscan", "testssl"],
        "password": ["john", "hashcat", "hashid", "hash-identifier", "crunch", "cewl", "hydra", "medusa"],
        "forensics": ["exiftool", "foremost", "binwalk", "bulk_extractor", "scalpel", "testdisk", "photorec", "fls", "icat", "istat"],
        "memory": ["volatility", "vol.py", "vol3"],
        "malware": ["yara", "ssdeep", "peepdf", "pdfid", "pdf-parser", "pdftotext", "olevba", "oleid", "oledir", "peframe", "pestudio", "floss", "capa"],
        "android_re": ["apktool", "jadx", "dex2jar", "apksigner", "aapt", "aapt2", "zipgrep"],
        "binary_re": ["radare2", "r2", "rabin2", "rasm2", "rax2", "rahash2", "ghidra", "readelf", "objdump", "nm", "ltrace", "strace", "gdb"],
        "rootkit_scan": ["rkhunter", "chkrootkit"],
        "antivirus": ["clamscan", "clamav", "sigtool"],
        "fake_services_remnux": ["inetsim", "fakedns", "dnschef"],
        "url_recon": ["gau", "waybackurls", "katana", "hakrawler", "httprobe", "subjs", "tlsx"],
        "data_tools": ["jq", "gron", "zipgrep", "yextend"],
    }
    out = {}
    for cat, tools in categories.items():
        installed = [t for t in tools if shutil.which(t)]
        missing = [t for t in tools if not shutil.which(t)]
        out[cat] = {"installed": installed, "missing": missing}
    return out


# ===========================================================================
# PDF Forensics / Redaction Analysis Tools
# ===========================================================================

@mcp.tool()
def pdf_forensic_analysis(pdf_path: str) -> dict:
    """Deep forensic analysis of a PDF: revision history, metadata, software signatures,
    page consistency, embedded objects, and cross-reference structure."""
    import fitz
    p = _require_path(pdf_path)
    doc = fitz.open(str(p))

    with open(str(p), "rb") as f:
        content = f.read()

    eof_count = content.count(b"%%EOF")
    eof_positions = [m.start() for m in re.finditer(b"%%EOF", content)]

    signatures = {}
    for sig, name in [(b"Adobe", "Adobe"), (b"Acrobat", "Acrobat"), (b"Phantom", "PhantomPDF"),
                      (b"ABBYY", "ABBYY_OCR"), (b"Tesseract", "Tesseract"), (b"Relativity", "Relativity"),
                      (b"Nuance", "Nuance"), (b"Kofax", "Kofax")]:
        if sig in content:
            idx = content.find(sig)
            signatures[name] = content[max(0,idx-10):idx+40].decode("latin-1", errors="replace")

    page_info = []
    for i in range(min(10, len(doc))):
        page = doc[i]
        images = page.get_images(full=True)
        text_len = len(page.get_text("text"))
        page_info.append({"page": i+1, "width": round(page.rect.width), "height": round(page.rect.height),
                         "images": len(images), "text_len": text_len})

    meta = doc.metadata
    result = {
        "path": str(p), "pages": len(doc), "file_size": len(content),
        "pdf_version": meta.get("format", ""), "creator": meta.get("creator", ""),
        "producer": meta.get("producer", ""), "creation_date": meta.get("creationDate", ""),
        "mod_date": meta.get("modDate", ""), "revisions": eof_count,
        "eof_positions": eof_positions, "software_signatures": signatures,
        "page_analysis": page_info,
        "unique_page_sizes": len(set((p["width"], p["height"]) for p in page_info)),
    }
    doc.close()
    return result


@mcp.tool()
def pdf_extract_revision(pdf_path: str, revision: int = 1) -> dict:
    """Extract a specific revision from a PDF with incremental saves.
    Revision 1 = first version (potentially pre-redacted). Returns text from that revision."""
    import fitz
    p = _require_path(pdf_path)

    with open(str(p), "rb") as f:
        content = f.read()

    eof_positions = [m.start() + 5 for m in re.finditer(b"%%EOF", content)]
    if revision > len(eof_positions):
        return {"error": f"Only {len(eof_positions)} revisions found"}

    rev_content = content[:eof_positions[revision - 1] + 1]
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False, dir=str(SCRATCH)) as tf:
        tf.write(rev_content)
        rev_path = tf.name

    doc = fitz.open(rev_path)
    text = "\n".join([doc[i].get_text("text") for i in range(len(doc))])
    pages = len(doc)
    doc.close()

    import os
    os.unlink(rev_path)
    return {"revision": revision, "pages": pages, "size": len(rev_content), "text_preview": text[:10000]}


@mcp.tool()
def pdf_deredact_crossref(redacted_path: str, unredacted_path: str) -> dict:
    """Cross-reference two versions of the same document to recover redacted content.
    Compares line-by-line and returns content only present in the less-redacted version."""
    import fitz
    p1 = _require_path(redacted_path)
    p2 = _require_path(unredacted_path)

    doc1 = fitz.open(str(p1))
    doc2 = fitz.open(str(p2))

    text1 = "\n".join([doc1[i].get_text("text") for i in range(len(doc1))])
    text2 = "\n".join([doc2[i].get_text("text") for i in range(len(doc2))])
    doc1.close()
    doc2.close()

    lines1 = set(l.strip() for l in text1.split("\n") if l.strip() and len(l.strip()) > 10)
    lines2 = set(l.strip() for l in text2.split("\n") if l.strip() and len(l.strip()) > 10)

    only_in_less_redacted = lines2 - lines1
    red_count_1 = text1.lower().count("redacted")
    red_count_2 = text2.lower().count("redacted")

    meaningful = [l for l in sorted(only_in_less_redacted) if not l.startswith("EFTA") and len(l) > 15]

    return {
        "redactions_in_doc1": red_count_1,
        "redactions_in_doc2": red_count_2,
        "unique_to_less_redacted": len(meaningful),
        "recovered_content": meaningful[:500],
    }


@mcp.tool()
def pdf_redaction_image_analysis(pdf_path: str, page_num: int, dpi: int = 600) -> dict:
    """Analyze redaction boxes on a specific page using image processing.
    Applies gamma correction, contrast stretching, inversion, and threshold techniques
    to attempt recovery of text under/around redaction boxes."""
    import fitz
    import numpy as np
    from PIL import Image, ImageEnhance
    import pytesseract
    import io

    p = _require_path(pdf_path)
    doc = fitz.open(str(p))
    page = doc[page_num - 1]

    pix = page.get_pixmap(dpi=dpi)
    img = Image.open(io.BytesIO(pix.tobytes("png")))
    arr = np.array(img.convert("L"))

    standard_text = page.get_text("text")
    doc.close()

    height, width = arr.shape
    row_dark_pct = np.mean(arr < 30, axis=1)

    in_box = False
    boxes = []
    start = 0
    for y in range(height):
        if row_dark_pct[y] > 0.15 and not in_box:
            start = y
            in_box = True
        elif row_dark_pct[y] <= 0.15 and in_box:
            if y - start > 15:
                boxes.append((start, y))
            in_box = False

    results = []
    for i, (y_start, y_end) in enumerate(boxes[:20]):
        region = arr[max(0,y_start-3):min(height,y_end+3), :]

        transforms = {
            "gamma_03": np.clip(((region / 255.0) ** 0.3) * 255, 0, 255).astype(np.uint8),
            "gamma_5": np.clip(((region / 255.0) ** 5.0) * 255, 0, 255).astype(np.uint8),
            "threshold_20_80": ((region > 20) & (region < 80)).astype(np.uint8) * 255,
            "stretch": np.clip((region.astype(float) - region.min()) / max(1, region.max() - region.min()) * 255, 0, 255).astype(np.uint8),
        }

        for name, enhanced in transforms.items():
            try:
                text = pytesseract.image_to_string(Image.fromarray(enhanced), config="--psm 7").strip()
                if text and len(text) > 2 and text.lower() not in standard_text.lower() and text != "[REDACTED]":
                    results.append({"box": i, "y_range": f"{y_start}-{y_end}", "technique": name, "text": text})
            except:
                pass

    return {
        "page": page_num, "dpi": dpi, "image_size": f"{width}x{height}",
        "redaction_boxes_found": len(boxes),
        "recovered_text": results,
    }


@mcp.tool()
def pdf_raw_stream_extract(pdf_path: str, page_num: int) -> dict:
    """Extract raw content stream from a PDF page, looking for hidden text operators
    (BT/ET blocks) that might contain text rendered invisible or under redaction boxes."""
    import fitz
    p = _require_path(pdf_path)
    doc = fitz.open(str(p))
    page = doc[page_num - 1]

    xref = page.xref
    raw = doc.xref_stream(xref)

    text_blocks = []
    if raw:
        import re as re_mod
        bt_blocks = re_mod.findall(rb"BT(.*?)ET", raw, re_mod.DOTALL)
        for block in bt_blocks[:50]:
            strings = re_mod.findall(rb"\((.*?)\)", block)
            decoded = []
            for s in strings:
                try:
                    decoded.append(s.decode("latin-1"))
                except:
                    pass
            if decoded:
                combined = "".join(decoded).strip()
                if combined and len(combined) > 1:
                    text_blocks.append(combined)

    visible_text = page.get_text("text")
    hidden = [t for t in text_blocks if t not in visible_text and len(t) > 2]

    doc.close()
    return {
        "page": page_num, "total_stream_text_blocks": len(text_blocks),
        "visible_text_length": len(visible_text),
        "potentially_hidden_text": hidden[:100],
    }


@mcp.tool()
def pdf_find_duplicates(directory: str, min_overlap: int = 8) -> dict:
    """Find potential duplicate document pairs in a directory that may have different
    redaction levels. Uses text fingerprinting to identify documents with similar content."""
    import fitz
    d = _require_path(directory)

    fingerprints = {}
    for pdf in sorted(d.glob("*.pdf"))[:1000]:
        try:
            doc = fitz.open(str(pdf))
            text = ""
            for page in doc:
                text += page.get_text("text")
                if len(text) > 500:
                    break
            doc.close()
            clean = re.sub(r"\[REDACTED\]|\(REDACTED\)|\{REDACTED\}", "XXX", text[:400])
            words = [w for w in clean.split() if len(w) > 4 and w != "XXX"][:20]
            fp = " ".join(words)
            if fp and len(fp) > 40:
                if fp not in fingerprints:
                    fingerprints[fp] = []
                fingerprints[fp].append(pdf.name)
        except:
            pass

    duplicates = {fp: docs for fp, docs in fingerprints.items() if len(docs) > 1}
    return {
        "directory": str(d), "documents_scanned": len(fingerprints),
        "duplicate_groups": len(duplicates),
        "groups": {", ".join(docs): fp[:100] for fp, docs in duplicates.items()},
    }


# ===========================================================================
# Document Research Tools
# ===========================================================================
# Tools for investigating large document collections: ZIP archives, PDFs,
# images (OCR), Office docs, spreadsheets, email, and full-text search.
# ===========================================================================

DOC_INDEX_DIR = SCRATCH / "doc-index"
DOC_INDEX_DIR.mkdir(parents=True, exist_ok=True)


@mcp.tool()
def archive_list(archive_path: str, grep_pattern: str | None = None, limit: int = 2000) -> dict:
    """List contents of a ZIP/tar archive without extracting. Optionally filter by regex."""
    import zipfile, tarfile
    p = _require_path(archive_path)
    entries = []
    if zipfile.is_zipfile(str(p)):
        with zipfile.ZipFile(str(p), "r") as zf:
            for info in zf.infolist():
                entries.append({"name": info.filename, "size": info.file_size,
                                "compressed": info.compress_size, "date": f"{info.date_time[0]:04d}-{info.date_time[1]:02d}-{info.date_time[2]:02d}"})
    elif tarfile.is_tarfile(str(p)):
        with tarfile.open(str(p), "r:*") as tf:
            for m in tf.getmembers():
                entries.append({"name": m.name, "size": m.size, "date": ""})
    else:
        return {"error": "Not a recognized ZIP or tar archive"}
    if grep_pattern:
        regex = re.compile(grep_pattern, re.IGNORECASE)
        entries = [e for e in entries if regex.search(e["name"])]
    total = len(entries)
    return {"archive": str(p), "total_entries": total, "entries": entries[:limit]}


@mcp.tool()
def archive_extract(archive_path: str, out_dir: str | None = None, members: list[str] | None = None) -> dict:
    """Extract files from a ZIP/tar archive. If members is set, only extract those paths."""
    import zipfile, tarfile
    p = _require_path(archive_path)
    if not out_dir:
        out_dir = str(SCRATCH / f"extract-{p.stem}")
    od = Path(out_dir).expanduser()
    od.mkdir(parents=True, exist_ok=True)
    extracted = []
    if zipfile.is_zipfile(str(p)):
        with zipfile.ZipFile(str(p), "r") as zf:
            targets = members if members else [m.filename for m in zf.infolist() if not m.is_dir()]
            for name in targets:
                try:
                    zf.extract(name, path=str(od))
                    extracted.append(name)
                except (KeyError, Exception) as e:
                    extracted.append(f"ERROR:{name}:{e}")
    elif tarfile.is_tarfile(str(p)):
        with tarfile.open(str(p), "r:*") as tf:
            targets = members if members else [m.name for m in tf.getmembers() if m.isfile()]
            for name in targets:
                try:
                    tf.extract(name, path=str(od), filter="data")
                    extracted.append(name)
                except Exception as e:
                    extracted.append(f"ERROR:{name}:{e}")
    return {"out_dir": str(od), "extracted_count": len(extracted), "files": extracted[:500]}


@mcp.tool()
def pdf_extract_pages(pdf_path: str, pages: str | None = None, ocr_fallback: bool = True) -> dict:
    """Extract text from a PDF with page-level granularity using PyMuPDF. Falls back to OCR for scanned pages.
    pages: e.g. '1-5', '3', '10-20'. None = all pages."""
    import fitz
    p = _require_path(pdf_path)
    doc = fitz.open(str(p))
    page_range = range(len(doc))
    if pages:
        parts = pages.split("-")
        if len(parts) == 2:
            page_range = range(max(0, int(parts[0]) - 1), min(len(doc), int(parts[1])))
        else:
            page_range = [int(parts[0]) - 1]
    results = []
    for i in page_range:
        page = doc[i]
        text = page.get_text("text").strip()
        method = "text"
        if not text and ocr_fallback:
            try:
                import pytesseract
                from PIL import Image
                import io
                pix = page.get_pixmap(dpi=300)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                text = pytesseract.image_to_string(img).strip()
                method = "ocr"
            except Exception as e:
                text = f"[OCR failed: {e}]"
                method = "ocr-failed"
        results.append({"page": i + 1, "method": method, "text": text[:10000]})
    doc.close()
    return {"path": str(p), "total_pages": len(doc) if hasattr(doc, '__len__') else "?", "pages": results}


@mcp.tool()
def ocr_image(image_path: str, lang: str = "eng", psm: int = 3) -> dict:
    """OCR an image file (PNG, JPG, TIFF, BMP) via Tesseract. psm: page segmentation mode (3=auto, 6=block, 11=sparse)."""
    import pytesseract
    from PIL import Image
    p = _require_path(image_path)
    img = Image.open(str(p))
    config = f"--psm {psm}"
    text = pytesseract.image_to_string(img, lang=lang, config=config)
    return {"path": str(p), "size": f"{img.width}x{img.height}", "lang": lang, "text": _truncate(text, 50000)}


@mcp.tool()
def docx_extract(docx_path: str) -> dict:
    """Extract text, tables, and metadata from a Word .docx file."""
    from docx import Document
    p = _require_path(docx_path)
    doc = Document(str(p))
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    tables_out = []
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables_out.append({"table_index": i, "rows": rows[:100]})
    props = doc.core_properties
    meta = {
        "author": props.author, "title": props.title, "subject": props.subject,
        "created": str(props.created), "modified": str(props.modified),
        "last_modified_by": props.last_modified_by,
    }
    return {"path": str(p), "metadata": meta, "paragraphs": paragraphs[:500],
            "table_count": len(tables_out), "tables": tables_out[:20]}


@mcp.tool()
def xlsx_extract(xlsx_path: str, sheet_name: str | None = None, limit: int = 500) -> dict:
    """Read an Excel .xlsx file. Returns sheet names, headers, and rows as dicts."""
    from openpyxl import load_workbook
    p = _require_path(xlsx_path)
    wb = load_workbook(str(p), read_only=True, data_only=True)
    sheets = wb.sheetnames
    target = sheet_name or sheets[0]
    ws = wb[target]
    rows = []
    headers = None
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        if i == 0:
            headers = [str(c) if c else f"col_{j}" for j, c in enumerate(row)]
            continue
        if i > limit:
            break
        rows.append({h: (str(v) if v is not None else "") for h, v in zip(headers, row)})
    wb.close()
    return {"path": str(p), "sheets": sheets, "active_sheet": target,
            "headers": headers, "row_count": len(rows), "rows": rows}


@mcp.tool()
def email_parse(eml_path: str) -> dict:
    """Parse a .eml email file: headers, body text, and attachment list."""
    import email
    from email import policy
    p = _require_path(eml_path)
    with open(str(p), "rb") as f:
        msg = email.message_from_binary_file(f, policy=policy.default)
    headers = {k: msg[k] for k in ["From", "To", "Cc", "Bcc", "Subject", "Date", "Message-ID"] if msg[k]}
    body_parts = []
    attachments = []
    for part in msg.walk():
        ct = part.get_content_type()
        fn = part.get_filename()
        if fn:
            attachments.append({"filename": fn, "content_type": ct, "size": len(part.get_payload(decode=True) or b"")})
        elif ct in ("text/plain", "text/html"):
            payload = part.get_payload(decode=True)
            if payload:
                body_parts.append({"type": ct, "text": payload.decode("utf-8", errors="replace")[:20000]})
    return {"path": str(p), "headers": headers, "body": body_parts, "attachments": attachments}


@mcp.tool()
def doc_search(directory: str, query: str, extensions: list[str] | None = None, limit: int = 100) -> dict:
    """Full-text search across documents in a directory. Searches PDF text, docx, txt, csv, and filenames.
    Returns matching files with context snippets."""
    d = _require_path(directory)
    exts = extensions or [".pdf", ".txt", ".csv", ".docx", ".doc", ".eml", ".htm", ".html", ".json", ".xml", ".md", ".log"]
    query_lower = query.lower()
    matches = []
    for p in d.rglob("*"):
        if not p.is_file():
            continue
        if p.suffix.lower() not in exts and query_lower not in p.name.lower():
            continue
        if query_lower in p.name.lower():
            matches.append({"file": str(p.relative_to(d)), "match_type": "filename", "snippet": p.name})
        try:
            if p.suffix.lower() == ".pdf":
                import fitz
                doc = fitz.open(str(p))
                for i in range(min(len(doc), 50)):
                    text = doc[i].get_text("text")
                    if query_lower in text.lower():
                        idx = text.lower().index(query_lower)
                        snippet = text[max(0, idx - 100):idx + len(query) + 100].strip()
                        matches.append({"file": str(p.relative_to(d)), "match_type": "pdf_content", "page": i + 1, "snippet": snippet})
                doc.close()
            elif p.suffix.lower() == ".docx":
                from docx import Document
                doc = Document(str(p))
                for para in doc.paragraphs:
                    if query_lower in para.text.lower():
                        matches.append({"file": str(p.relative_to(d)), "match_type": "docx_content", "snippet": para.text[:200]})
                        break
            else:
                text = p.read_text(errors="replace")[:500000]
                if query_lower in text.lower():
                    idx = text.lower().index(query_lower)
                    snippet = text[max(0, idx - 100):idx + len(query) + 100].strip()
                    matches.append({"file": str(p.relative_to(d)), "match_type": "text_content", "snippet": snippet})
        except Exception:
            pass
        if len(matches) >= limit:
            break
    return {"directory": str(d), "query": query, "match_count": len(matches), "matches": matches}


@mcp.tool()
def extract_entities(text_or_path: str, patterns: list[str] | None = None) -> dict:
    """Extract structured entities from text or a file: emails, phone numbers, dates, URLs, names (proper nouns), SSNs, dollar amounts.
    Pass a file path (must exist) or raw text."""
    p = Path(text_or_path).expanduser().resolve()
    if p.exists() and p.is_file():
        if p.suffix.lower() == ".pdf":
            import fitz
            doc = fitz.open(str(p))
            text = "\n".join(doc[i].get_text("text") for i in range(min(len(doc), 30)))
            doc.close()
        else:
            text = p.read_text(errors="replace")[:500000]
    else:
        text = text_or_path
    builtin = {
        "emails": r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+",
        "phones": r"\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}",
        "dates": r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4}\b",
        "urls": r"https?://[A-Za-z0-9._~:/?#\[\]@!$&'()*+,;=%-]+",
        "ssns": r"\b\d{3}-\d{2}-\d{4}\b",
        "dollar_amounts": r"\$[\d,]+(?:\.\d{2})?",
    }
    if patterns:
        for i, pat in enumerate(patterns):
            builtin[f"custom_{i}"] = pat
    results = {}
    for name, pat in builtin.items():
        found = list(set(re.findall(pat, text)))
        if found:
            results[name] = sorted(found)[:200]
    return {"source": str(p) if p.exists() else "raw_text", "entities": results}


@mcp.tool()
def doc_timeline(directory: str, extensions: list[str] | None = None, limit: int = 500) -> dict:
    """Build a timeline from file metadata (created/modified dates) and dates found inside documents."""
    d = _require_path(directory)
    exts = extensions or [".pdf", ".docx", ".doc", ".txt", ".eml", ".xlsx", ".jpg", ".png"]
    events = []
    for p in d.rglob("*"):
        if not p.is_file() or p.suffix.lower() not in exts:
            continue
        stat = p.stat()
        events.append({"file": str(p.relative_to(d)), "event": "modified",
                        "timestamp": int(stat.st_mtime), "date": __import__("datetime").datetime.fromtimestamp(stat.st_mtime).isoformat()})
        if len(events) >= limit:
            break
    events.sort(key=lambda e: e["timestamp"])
    return {"directory": str(d), "event_count": len(events), "timeline": events}


@mcp.tool()
def batch_pdf_text(directory: str, out_dir: str | None = None, ocr_if_empty: bool = True, limit: int = 200) -> dict:
    """Batch-extract text from all PDFs in a directory. Saves each as a .txt file in out_dir."""
    import fitz
    d = _require_path(directory)
    if not out_dir:
        out_dir = str(SCRATCH / f"pdf-text-{d.name}")
    od = Path(out_dir).expanduser()
    od.mkdir(parents=True, exist_ok=True)
    processed = []
    for i, pdf in enumerate(sorted(d.rglob("*.pdf"))):
        if i >= limit:
            break
        try:
            doc = fitz.open(str(pdf))
            text_parts = []
            for page in doc:
                t = page.get_text("text")
                if not t.strip() and ocr_if_empty:
                    try:
                        import pytesseract
                        from PIL import Image
                        import io
                        pix = page.get_pixmap(dpi=200)
                        img = Image.open(io.BytesIO(pix.tobytes("png")))
                        t = pytesseract.image_to_string(img)
                    except Exception:
                        pass
                text_parts.append(t)
            doc.close()
            full_text = "\n\n--- PAGE BREAK ---\n\n".join(text_parts)
            rel = pdf.relative_to(d)
            out_file = od / f"{rel.stem}.txt"
            out_file.parent.mkdir(parents=True, exist_ok=True)
            out_file.write_text(full_text, errors="replace")
            processed.append({"pdf": str(rel), "pages": len(text_parts), "chars": len(full_text), "out": str(out_file)})
        except Exception as e:
            processed.append({"pdf": str(pdf.relative_to(d)), "error": str(e)})
    return {"source_dir": str(d), "out_dir": str(od), "processed": len(processed), "files": processed}


@mcp.tool()
def doc_index_build(directory: str, index_name: str = "default") -> dict:
    """Build a full-text search index (Whoosh) over documents in a directory. Indexes PDF, txt, docx, eml, csv."""
    from whoosh.index import create_in
    from whoosh.fields import Schema, TEXT, ID, STORED
    d = _require_path(directory)
    idx_dir = DOC_INDEX_DIR / index_name
    idx_dir.mkdir(parents=True, exist_ok=True)
    schema = Schema(path=ID(stored=True, unique=True), filename=TEXT(stored=True), content=TEXT(stored=True))
    ix = create_in(str(idx_dir), schema)
    writer = ix.writer()
    count = 0
    for p in d.rglob("*"):
        if not p.is_file():
            continue
        text = ""
        try:
            if p.suffix.lower() == ".pdf":
                import fitz
                doc = fitz.open(str(p))
                text = "\n".join(doc[i].get_text("text") for i in range(min(len(doc), 100)))
                doc.close()
            elif p.suffix.lower() == ".docx":
                from docx import Document
                doc = Document(str(p))
                text = "\n".join(para.text for para in doc.paragraphs)
            elif p.suffix.lower() in (".txt", ".csv", ".eml", ".htm", ".html", ".json", ".xml", ".md", ".log"):
                text = p.read_text(errors="replace")[:500000]
        except Exception:
            continue
        if text.strip():
            writer.add_document(path=str(p), filename=p.name, content=text[:200000])
            count += 1
    writer.commit()
    return {"index": index_name, "index_dir": str(idx_dir), "documents_indexed": count}


@mcp.tool()
def doc_index_search(query: str, index_name: str = "default", limit: int = 50) -> dict:
    """Search the full-text index built by doc_index_build. Returns matching documents with snippets."""
    from whoosh.index import open_dir
    from whoosh.qparser import QueryParser
    idx_dir = DOC_INDEX_DIR / index_name
    if not idx_dir.exists():
        return {"error": f"Index '{index_name}' not found. Run doc_index_build first."}
    ix = open_dir(str(idx_dir))
    with ix.searcher() as searcher:
        qp = QueryParser("content", ix.schema)
        q = qp.parse(query)
        results = searcher.search(q, limit=limit)
        hits = []
        for hit in results:
            content = hit.get("content", "")
            ql = query.lower()
            idx = content.lower().find(ql)
            snippet = content[max(0, idx - 150):idx + len(query) + 150].strip() if idx >= 0 else content[:300]
            hits.append({"path": hit["path"], "filename": hit["filename"], "score": round(hit.score, 3), "snippet": snippet})
    return {"query": query, "index": index_name, "hit_count": len(hits), "hits": hits}


# ----- Server entry point ----------------------------------------------------

def main():
    mcp.run()


if __name__ == "__main__":
    main()
